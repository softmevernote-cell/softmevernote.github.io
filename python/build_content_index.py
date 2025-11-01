#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_content_index_v10_emit_parts.py

캐시 정책 (요청 반영):
- --emit name  : HTML 캐시를 **항상 재생성**(이미 있으면 삭제/덮어쓰기)
- --emit html  : HTML 캐시를 **없을 때만 생성**
- --emit attach: HTML 캐시를 **없을 때만 생성** + 첨부 캐시는 **항상 재생성**

공통 스키마
- name  -> content_index_for_name.json
  (html_file, folder, keywords_name, summary_name, tags, subtags, date, source, html_text_ref)
- html  -> content_index_for_html.json
  (html_file, folder, keywords_html, summary_html, tags, subtags, date, source, html_text_ref)
- attach-> content_index_for_attach.json
  (html_file, folder, keywords_attach, summary_file, attachments[], tags, subtags, date, source, html_text_ref)
"""
import argparse, json, re, unicodedata as ud, io, zipfile, html
from collections import Counter
from pathlib import Path
from typing import Dict, List, Any, Optional

# ===== Optional deps detect =====
BS4 = None
PDF_EXTRACT = None
DOCX_EXTRACT = None
PIL = None
TESS = None
OPENPYXL = None
LXML_ETREE = None
PYHWP = None

try:
    from bs4 import BeautifulSoup as _BS4, Comment as _Comment
    BS4 = _BS4
    BS4_Comment = _Comment
except Exception:
    BS4 = None
    BS4_Comment = None

try:
    from pdfminer.high_level import extract_text as _pdf_extract_text
    PDF_EXTRACT = _pdf_extract_text
except Exception:
    PDF_EXTRACT = None

try:
    import docx as _docx
    DOCX_EXTRACT = _docx
except Exception:
    DOCX_EXTRACT = None

try:
    from PIL import Image as _Image
    PIL = _Image
    import pytesseract as _tess
    TESS = _tess
except Exception:
    PIL = None
    TESS = None

try:
    import openpyxl as _openpyxl
    OPENPYXL = _openpyxl
except Exception:
    OPENPYXL = None

try:
    from lxml import etree as _etree
    LXML_ETREE = _etree
except Exception:
    LXML_ETREE = None

try:
    import pyhwp as _pyhwp
    PYHWP = _pyhwp
except Exception:
    PYHWP = None

DEFAULT_STOPWORDS = set("""
the and or of to a an in on for at by with from as is are was were be been being
및 그리고 등 이 가 은 는 을 를 에 의 와 과 이다 하다 있다 없는 대한 관련 자료 문서 파일 내가 너가 내 노트
""".split())

IMAGE_EXTS = {"jpg","jpeg","png","gif","bmp","tiff","webp"}
TEXTLIKE_EXTS = {"txt","csv","tsv","md","log","json"}
SHEET_EXTS = {"xlsx","xls","csv"}
DOC_EXTS = {"docx","doc"}
PDF_EXTS = {"pdf"}
MHTML_EXTS = {"mht","mhtml"}
ARCHIVE_EXTS = {"zip"}
HANGUL_EXTS = {"hwp","hwpx"}

def nfc(s: str) -> str:
    return ud.normalize("NFC", s or "")

def slugify_path_component(s: str) -> str:
    s = nfc(s)
    s = re.sub(r"[\\/:*?\"<>|]+", "_", s)
    s = re.sub(r"\s+", "_", s).strip("_")
    return s[:200] if len(s) > 200 else s

def read_json(path: Path) -> Any:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def write_text(path: Path, text: str, overwrite: bool=False):
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists() and not overwrite:
        return False
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return True

def write_json(path: Path, obj: Any):
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)

def tokenize(s: str, lowercase=True, strip_years=True, stopwords: Optional[set]=None) -> List[str]:
    if not s: return []
    s = nfc(s)
    if lowercase:
        s = s.lower()
    s = re.sub(r"[\[\]{}()<>:;,.!?\"'`~@#$%^&*+=\\\/|_]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    toks = [t for t in s.split(" ") if t]
    if strip_years:
        toks = [t for t in toks if not re.fullmatch(r"(19|20)\d{2}", t)]
    if stopwords:
        toks = [t for t in toks if t not in stopwords and len(t) >= 2]
    return toks

def decode_html_entities_and_remove_some(s: str) -> str:
    if not s:
        return s
    for ent in ("&nbsp;", "&quot;", "&lt;"):
        s = s.replace(ent, "")
    s = html.unescape(s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def extract_text_from_html_file(path: Path) -> str:
    data = None
    try:
        data = path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        try:
            data = path.read_text(encoding="cp949", errors="ignore")
        except Exception:
            return ""
    if data is None:
        return ""

    if BS4:
        try:
            soup = _BS4(data, "html.parser")
            for el in soup(["script","style","noscript","meta","link","head","title","template"]):
                el.decompose()
            if BS4_Comment:
                for c in soup.find_all(string=lambda t: isinstance(t, BS4_Comment)):
                    c.extract()
            for el in soup.select("[hidden], input[type='hidden'], [aria-hidden]"):
                el.decompose()
            for el in soup.select("[style]"):
                style = (el.get("style") or "").lower().replace(" ","")
                if any(k in style for k in ("display:none","visibility:hidden","opacity:0","font-size:0")):
                    el.decompose()
            HIDDEN_CLASSES = {"hidden","d-none","sr-only","visually-hidden","screen-reader-text","a11y-hidden","u-hidden","is-hidden"}
            for el in soup.find_all(True, class_=True):
                classes = set(str(c).strip().lower() for c in (el.get("class") or []))
                if classes & HIDDEN_CLASSES:
                    el.decompose()
            root = soup.body if soup.body else soup
            text = root.get_text(separator=" ", strip=True)
            text = re.sub(r"\s+", " ", text)
            return decode_html_entities_and_remove_some(nfc(text.strip()))
        except Exception:
            pass

    s = data
    m = re.search(r"(?is)<body[^>]*>(.*?)</body>", s)
    if m:
        s = m.group(1)
    s = re.sub(r"(?is)<!--.*?-->", " ", s)
    s = re.sub(r"(?is)<(script|style|template|noscript|head|title)[^>]*>.*?</\1>", " ", s)
    s = re.sub(r"(?is)<(meta|link)(\s[^>]*?)?>", " ", s)
    s = re.sub(r"(?is)<[^>]+>", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return decode_html_entities_and_remove_some(s)

# --- Attachments extractors ---
def extract_text_from_pdf(path: Path) -> str:
    if PDF_EXTRACT is None:
        return ""
    try:
        txt = PDF_EXTRACT(str(path)) or ""
        return nfc(re.sub(r"\s+", " ", txt).strip())
    except Exception:
        return ""

def extract_text_from_docx(path: Path) -> str:
    if DOCX_EXTRACT is None:
        return ""
    try:
        doc = DOCX_EXTRACT.Document(str(path))
        parts = []
        for p in doc.paragraphs:
            if p.text: parts.append(p.text)
        return nfc(re.sub(r"\s+", " ", " ".join(parts)).strip())
    except Exception:
        return ""

def extract_text_from_txt(path: Path) -> str:
    try:
        data = path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        try:
            data = path.read_text(encoding="cp949", errors="ignore")
        except Exception:
            data = ""
    return nfc(re.sub(r"\s+", " ", data).strip())

def extract_text_from_image(path: Path, ocr: bool=False) -> str:
    if not ocr or PIL is None or TESS is None:
        return ""
    try:
        img = PIL.open(str(path))
        txt = TESS.image_to_string(img, lang="kor+eng")
        return nfc(re.sub(r"\s+", " ", txt).strip())
    except Exception:
        return ""

def extract_text_from_xlsx(path: Path) -> str:
    if OPENPYXL is None:
        return ""
    try:
        wb = OPENPYXL.load_workbook(filename=str(path), data_only=True, read_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" ".join(cells))
        return nfc(re.sub(r"\s+", " ", " ".join(parts)).strip())
    except Exception:
        return ""

def extract_text_from_hwpx(path: Path) -> str:
    if LXML_ETREE is None:
        return ""
    try:
        out = []
        with zipfile.ZipFile(str(path), 'r') as zf:
            for name in zf.namelist():
                if name.lower().endswith((".xml",)):
                    try:
                        data = zf.read(name)
                        root = LXML_ETREE.fromstring(data)
                        texts = root.xpath("//text()")
                        if texts:
                            out.append(" ".join([t for t in texts if isinstance(t, str)]))
                    except Exception:
                        continue
        return nfc(re.sub(r"\s+", " ", " ".join(out)).strip())
    except Exception:
        return ""

def extract_text_from_hwp(path: Path) -> str:
    if PYHWP is None:
        return ""
    try:
        from io import StringIO
        buf = StringIO()
        doc = PYHWP.HWPDocument(str(path))
        for sec in doc.bodytext.section_list:
            for p in sec.paragraph_list:
                txt = p.get_text()
                if txt:
                    buf.write(txt + " ")
        return nfc(re.sub(r"\s+", " ", buf.getvalue()).strip())
    except Exception:
        return ""

def extract_text_from_zip(path: Path, ocr: bool=False) -> str:
    try:
        out = []
        with zipfile.ZipFile(str(path), 'r') as zf:
            for name in zf.namelist():
                if name.endswith("/"):
                    continue
                ext = Path(name).suffix.lower().lstrip(".")
                try:
                    data = zf.read(name)
                except Exception:
                    continue
                tmp = Path(f"/tmp/{slugify_path_component(path.stem)}__{slugify_path_component(name)}")
                try:
                    tmp.parent.mkdir(parents=True, exist_ok=True)
                    with open(tmp, "wb") as f:
                        f.write(data)
                except Exception:
                    continue

                if ext in PDF_EXTS:
                    out.append(extract_text_from_pdf(tmp))
                elif ext in {"docx","doc"}:
                    out.append(extract_text_from_docx(tmp))
                elif ext in {"txt","csv","tsv","md","log","json"}:
                    out.append(extract_text_from_txt(tmp))
                elif ext in {"xlsx"}:
                    out.append(extract_text_from_xlsx(tmp))
                elif ext in {"hwpx"}:
                    out.append(extract_text_from_hwpx(tmp))
                elif ext in {"hwp"}:
                    out.append(extract_text_from_hwp(tmp))
                elif ext in {"html","htm"}:
                    out.append(extract_text_from_html_file(tmp))
                elif ext in IMAGE_EXTS:
                    out.append(extract_text_from_image(tmp, ocr=ocr))

                try: tmp.unlink(missing_ok=True)
                except Exception: pass

        return nfc(re.sub(r"\s+", " ", " ".join([t for t in out if t])).strip())
    except Exception:
        return ""

def guess_type(name: str) -> str:
    ext = (Path(name).suffix.lower().lstrip(".")) if name else ""
    if ext in PDF_EXTS: return "pdf"
    if ext in DOC_EXTS: return "docx" if ext == "docx" else "doc"
    if ext in IMAGE_EXTS: return "image"
    if ext in TEXTLIKE_EXTS: return "text"
    if ext in MHTML_EXTS: return "mhtml"
    if ext in SHEET_EXTS: return "sheet"
    if ext in ARCHIVE_EXTS: return "zip"
    if ext in HANGUL_EXTS: return ext
    if ext in {"html","htm"}: return "html"
    return ext or "file"

def summarize(text: str, max_chars: int = 240) -> str:
    if not text:
        return ""
    t = re.split(r"(?<=[.!?다요])\s+", text)
    out = ""
    for s in t:
        if len(out) + len(s) + 1 > max_chars:
            break
        out = (out + " " + s).strip()
        if len(out) >= max_chars * 0.6:
            break
    return out[:max_chars].strip()

def top_keywords(text, limit, stopwords):
    cnt = Counter(tokenize(text, lowercase=True, strip_years=True, stopwords=stopwords))
    return [w for w,_ in cnt.most_common(limit)]

def process_record(rec: Dict[str, Any], root: Path, out_text_dir: Path,
                   k_name:int, k_html:int, k_attach:int,
                   stopwords:set, use_ocr: bool, only_attachments: bool,
                   emit_set:set):
    html_file_field = nfc(rec.get("html_file",""))
    folder_label = nfc(rec.get("folder",""))
    files = rec.get("files",[]) or []

    html_rel = Path(html_file_field + ".html")
    html_path = (root / html_rel).resolve()

    # ---- HTML cache path
    html_cache_name = f"{slugify_path_component(html_file_field)}_html.txt"
    html_cache_path = out_text_dir / html_cache_name
    html_text_ref = str(Path("text_cache") / html_cache_name)

    html_text = ""
    if html_path.exists():
        # Policy by emit
        if "name" in emit_set:
            # overwrite always
            try:
                html_text = extract_text_from_html_file(html_path)
            except Exception:
                html_text = ""
            write_text(html_cache_path, html_text, overwrite=True)
        elif "html" in emit_set or "attach" in emit_set:
            # create if missing
            if not html_cache_path.exists():
                try:
                    html_text = extract_text_from_html_file(html_path)
                except Exception:
                    html_text = ""
                write_text(html_cache_path, html_text, overwrite=False)

        # if we need html keywords/summary, read from cache (or extract if still empty)
        if "html" in emit_set:
            if html_cache_path.exists():
                try:
                    html_text = html_cache_path.read_text(encoding="utf-8", errors="ignore")
                except Exception:
                    html_text = ""
            if not html_text:
                try:
                    html_text = extract_text_from_html_file(html_path)
                except Exception:
                    html_text = ""
                write_text(html_cache_path, html_text, overwrite=True)
    # else: no html file

    # filename keywords
    name_tokens_src = Path(html_rel).stem
    keywords_name = top_keywords(name_tokens_src, k_name, stopwords)
    summary_name = " ".join(keywords_name[:12])

    # html keywords/summary only for html emit
    keywords_html = top_keywords(html_text, k_html, stopwords) if ("html" in emit_set and html_text) else []
    summary_html = summarize(html_text) if ("html" in emit_set) else ""

    # attachments (only if 'attach' in emit)
    attach_infos = []
    keywords_attach = []
    summary_file = ""
    if "attach" in emit_set:
        attach_text_concat = []
        attach_dir = html_path.with_name(html_path.stem + " files")
        for fn in files:
            ftype = guess_type(fn)
            text = ""
            raw_path = attach_dir / fn
            cache_name = f"{slugify_path_component(html_file_field)}__{slugify_path_component(fn)}_att.txt"
            cache_path = out_text_dir / cache_name
            cache_ref = str(Path("text_cache") / cache_name)

            if raw_path.exists():
                # ALWAYS refresh attachment cache
                if ftype == "pdf":
                    text = extract_text_from_pdf(raw_path)
                elif ftype in {"docx","doc"}:
                    text = extract_text_from_docx(raw_path)
                elif ftype == "text":
                    text = extract_text_from_txt(raw_path)
                elif ftype == "image":
                    text = extract_text_from_image(raw_path, ocr=use_ocr)
                elif ftype == "sheet":
                    if raw_path.suffix.lower() == ".xlsx":
                        text = extract_text_from_xlsx(raw_path)
                    else:
                        text = extract_text_from_txt(raw_path)
                elif ftype == "zip":
                    text = extract_text_from_zip(raw_path, ocr=use_ocr)
                elif ftype == "hwpx":
                    text = extract_text_from_hwpx(raw_path)
                elif ftype == "hwp":
                    text = extract_text_from_hwp(raw_path)
                elif ftype == "html":
                    if not only_attachments:
                        text = extract_text_from_html_file(raw_path)
                    else:
                        text = ""

                write_text(cache_path, text or "", overwrite=True)
                text_ref = cache_ref
            else:
                text_ref = None

            attach_text_concat.append(text or "")
            attach_infos.append({
                "filename": fn,
                "type": ftype,
                "keywords": top_keywords(text, 20, stopwords) if text else [],
                "text_ref": text_ref
            })

        attach_all_text = " ".join(attach_text_concat)
        keywords_attach = top_keywords(attach_all_text, k_attach, stopwords) if attach_all_text else []
        summary_file = summarize(attach_all_text)

    # tags/subtags/date
    tags = []
    RULES = [
        (r"(Android|안드로이드|Webview|AsyncTask|Eclipse|Xdebug|Gradle|JQuery|jQuery|JavaScript|PowerShell|WSH|NAVER Tech Talk|디버깅)", "tech/dev"),
        (r"(연서|가족|어린이집|증명사진|아기새|여권|가족)", "life/family"),
        (r"(주차단속|집매매|철수확인서|인수인계|계약서|의견진술|익명신고|윤리경영|공정위|1종 보통 적성검사)", "legal/admin"),
        (r"(비염|안약|헤르페스|건강)", "health"),
        (r"(아이디어|계획|TODO|솔루션|관리툴|아이템)", "idea/project"),
        (r"(정약용|주52시간|뉴스|철학|사회|매일경제|Chosunbiz)", "society/thought"),
        (r"(NAS|시놀로지|DLNA|랜선|UTP|FTP|케이블)", "hardware/it"),
        (r"(영농|농업|스마트팜|논|밭|수확|파종|유통)", "agri/farm"),
    ]
    name_for_tag = Path(html_rel).name
    for pat, tag in RULES:
        if re.search(pat, name_for_tag, flags=re.IGNORECASE):
            tags.append(tag)
    tags = sorted(list(set(tags)))

    exts = set(Path(f).suffix.lower().lstrip(".") for f in files)
    subtags = []
    if "pdf" in exts: subtags.append("has_pdf")
    if "hwp" in exts or "hwpx" in exts: subtags.append("has_hanword")
    if "docx" in exts or "doc" in exts: subtags.append("has_docx")
    if exts.intersection({"jpg","jpeg","png","gif"}): subtags.append("has_images")
    if exts.intersection({"zip","rar"}): subtags.append("has_archives")
    if exts.intersection({"mht","mhtml"}): subtags.append("has_mht")
    if exts.intersection({"wma","mp3","wav"}): subtags.append("has_audio")
    if exts.intersection({"xlsx","xls","csv"}): subtags.append("has_spreadsheet")

    m = re.search(r"\b(19|20)\d{2}-\d{2}-\d{2}\b", str(html_rel))
    date = m.group(0) if m else None
    if not date:
        y = re.search(r"\b(19|20)\d{2}\b", str(html_rel))
        date = y.group(0) if y else None

    source = {
        "html_path": str(html_rel).replace("\\","/"),
        "files_dir": str(Path(html_rel).with_suffix("").as_posix() + " files/")
    }

    rec_name = {
        "html_file": html_file_field,
        "folder": folder_label,
        "keywords_name": keywords_name,
        "summary_name": summary_name,
        "tags": tags,
        "subtags": subtags,
        "date": date,
        "source": source,
        "html_text_ref": (html_text_ref if html_path.exists() else None),
    }

    rec_html = {
        "html_file": html_file_field,
        "folder": folder_label,
        "keywords_html": (keywords_html if "html" in emit_set else []),
        "summary_html": (summary_html if "html" in emit_set else ""),
        "tags": tags,
        "subtags": subtags,
        "date": date,
        "source": source,
        "html_text_ref": (html_text_ref if html_path.exists() else None),
    }

    rec_attach = {
        "html_file": html_file_field,
        "folder": folder_label,
        "keywords_attach": (keywords_attach if "attach" in emit_set else []),
        "summary_file": (summary_file if "attach" in emit_set else ""),
        "attachments": (attach_infos if "attach" in emit_set else []),
        "tags": tags,
        "subtags": subtags,
        "date": date,
        "source": source,
        "html_text_ref": (html_text_ref if html_path.exists() else None),
    }

    return rec_name, rec_html, rec_attach

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--root", required=True)
    ap.add_argument("--files", required=True)
    ap.add_argument("--outdir", default="data")
    ap.add_argument("--k-name", type=int, default=20)
    ap.add_argument("--k-html", type=int, default=50)
    ap.add_argument("--k-attach", type=int, default=40)
    ap.add_argument("--stopwords", default=None)
    ap.add_argument("--ocr", choices=["true","false"], default="false")
    ap.add_argument("--only-attachments", choices=["true","false"], default="false")
    ap.add_argument("--emit", default="name,html,attach", help="name,html,attach (복수 지정 가능)")
    ap.add_argument("--dry-run", action="store_true")
    args = ap.parse_args()

    # emit parse
    emit_raw = (args.emit or "name,html,attach").strip().lower()
    import re as _re
    emit_set = set([s.strip() for s in _re.split(r"[,+\\s]+", emit_raw) if s.strip()])
    invalid = emit_set - {"name","html","attach"}
    if invalid:
        raise SystemExit(f"[ERROR] --emit invalid values: {sorted(invalid)} (use name, html, attach)")

    root = Path(args.root).resolve()
    files_path = Path(args.files).resolve()
    outdir = Path(args.outdir).resolve()
    out_text_dir = outdir / "text_cache"

    if not files_path.exists():
        raise SystemExit(f"[ERROR] files_info.json not found: {files_path}")
    if not root.exists() or not root.is_dir():
        raise SystemExit(f"[ERROR] root not a directory: {root}")

    try:
        base = read_json(files_path)
    except Exception as e:
        raise SystemExit(f"[ERROR] failed to read files_info.json: {e}")

    stopwords = DEFAULT_STOPWORDS
    if args.stopwords:
        try:
            sw = read_json(Path(args.stopwords))
            stopwords = set(sw)
        except Exception as e:
            print(f"[WARN] stopwords read failed: {e}, using default.")

    use_ocr = args.ocr.lower() == "true"
    only_attachments = args.only_attachments.lower() == "true"

    name_rows, html_rows, attach_rows = [], [], []

    for rec in base:
        try:
            rn, rh, ra = process_record(
                rec=rec, root=root, out_text_dir=out_text_dir,
                k_name=args.k_name, k_html=args.k_html, k_attach=args.k_attach,
                stopwords=stopwords, use_ocr=use_ocr, only_attachments=only_attachments,
                emit_set=emit_set
            )
            if "name" in emit_set: name_rows.append(rn)
            if "html" in emit_set: html_rows.append(rh)
            if "attach" in emit_set: attach_rows.append(ra)
        except Exception as e:
            print(f"[WARN] record failed: {rec.get('html_file')} -> {e}")

    if args.dry_run:
        print(f"[DRY] total records: {len(base)}")
        if "name" in emit_set:
            print("[DRY] name preview:")
            print(json.dumps(name_rows[:2], ensure_ascii=False, indent=2))
        if "html" in emit_set:
            print("[DRY] html preview:")
            print(json.dumps(html_rows[:2], ensure_ascii=False, indent=2))
        if "attach" in emit_set:
            print("[DRY] attach preview:")
            print(json.dumps(attach_rows[:2], ensure_ascii=False, indent=2))
        return

    outdir.mkdir(parents=True, exist_ok=True)

    if "name" in emit_set:
        write_json(outdir / "content_index_for_name.json", name_rows)
        print(f"[WRITE] {outdir / 'content_index_for_name.json'} ({len(name_rows)} items)")
    if "html" in emit_set:
        write_json(outdir / "content_index_for_html.json", html_rows)
        print(f"[WRITE] {outdir / 'content_index_for_html.json'} ({len(html_rows)} items)")
    if "attach" in emit_set:
        write_json(outdir / "content_index_for_attach.json", attach_rows)
        print(f"[WRITE] {outdir / 'content_index_for_attach.json'} ({len(attach_rows)} items)")

    print("[DONE]")
    print(f"  - cache : {out_text_dir}")

if __name__ == "__main__":
    main()
